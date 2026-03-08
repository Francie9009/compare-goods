import streamlit as st
import zipfile
import os
import re
import difflib
from io import BytesIO


st.set_page_config(page_title="商品服務比對工具", page_icon="🔍", layout="wide")
st.title("🔍 商品/服務名稱比對工具")
st.caption("上傳檔案（PDF 或 Word）或直接貼上文字，自動比對商品服務名稱差異。")


# ══ 文字抽取 ══════════════════════════════════════════════════════════════════

def extract_text_from_pdf(file_bytes):
    import pdfplumber
    parts = []
    with pdfplumber.open(BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                parts.append(t)
    return "\n".join(parts)


def extract_text_from_docx(file_bytes):
    tmp_dir = "tmp_docx_cmp"
    if os.path.exists(tmp_dir):
        import shutil
        shutil.rmtree(tmp_dir)
    os.makedirs(tmp_dir)
    with zipfile.ZipFile(BytesIO(file_bytes)) as z:
        z.extractall(tmp_dir)
    xml_path = os.path.join(tmp_dir, "word", "document.xml")
    with open(xml_path, encoding="utf-8") as f:
        xml = f.read()
    texts = re.findall(r'<w:t[^>]*>([^<]+)</w:t>', xml)
    import shutil
    shutil.rmtree(tmp_dir)
    return " ".join(texts)


def extract_text(uploaded_file):
    ext = uploaded_file.name.lower().rsplit(".", 1)[-1]
    data = uploaded_file.read()
    uploaded_file.seek(0)
    if ext == "pdf":
        return extract_text_from_pdf(data)
    elif ext == "docx":
        return extract_text_from_docx(data)
    return data.decode("utf-8", errors="ignore")


# ══ 解析商品服務 ══════════════════════════════════════════════════════════════

NOISE_PREFIXES = [
    "intellectual property", "trademark", "batch", "official journal",
    "notification", "registrar", "page ", "pg.", "dec ", "jan ", "feb ",
    "mar ", "apr ", "may ", "jun ", "jul ", "aug ", "sep ", "oct ",
    "nov ", "class ",
]


def is_noise(s):
    sl = s.lower()
    if len(s) < 3:
        return True
    if re.match(r'^\d+$', s):
        return True
    if any(sl.startswith(n) for n in NOISE_PREFIXES):
        return True
    return False


def parse_items(body):
    body = re.sub(r'\s+', ' ', body).strip()
    raw = re.split(r'[;；]\s*|\n+', body)
    items = []
    for item in raw:
        item = item.strip().strip(".,，。 ")
        if item and not is_noise(item):
            items.append(item)
    return items


def parse_goods_services(text):
    result = {}
    class_pattern = re.compile(
        r'CLASS\s+(\d+)(.*?)(?=CLASS\s+\d+|\Z)',
        re.IGNORECASE | re.DOTALL
    )
    matches = list(class_pattern.finditer(text))
    if matches:
        for m in matches:
            cls = f"Class {m.group(1).strip()}"
            items = parse_items(m.group(2))
            if items:
                result[cls] = items
    else:
        items = parse_items(text)
        if items:
            result["（未分類）"] = items
    return result


# ══ 比對 ══════════════════════════════════════════════════════════════════════

def normalize(s):
    s = s.lower().strip()
    s = re.sub(r'\s+', ' ', s)
    s = s.replace('，', ',').replace('；', ';').replace('／', '/')
    return s


def compare_items(list_a, list_b):
    norm_a = {normalize(x): x for x in list_a}
    norm_b = {normalize(x): x for x in list_b}
    keys_a, keys_b = set(norm_a), set(norm_b)
    exact_same = keys_a & keys_b
    only_a = keys_a - keys_b
    only_b = keys_b - keys_a
    similar, matched_a, matched_b = [], set(), set()
    for ka in only_a:
        best_r, best_kb = 0, None
        for kb in only_b:
            r = difflib.SequenceMatcher(None, ka, kb).ratio()
            if r > best_r:
                best_r, best_kb = r, kb
        if best_r >= 0.70 and best_kb:
            similar.append((norm_a[ka], norm_b[best_kb], round(best_r, 2)))
            matched_a.add(ka)
            matched_b.add(best_kb)
    return {
        "same":    sorted([norm_a[k] for k in exact_same]),
        "only_a":  sorted([norm_a[k] for k in only_a if k not in matched_a]),
        "only_b":  sorted([norm_b[k] for k in only_b if k not in matched_b]),
        "similar": similar,
    }


def compare_all(parsed_a, parsed_b):
    all_cls = sorted(set(list(parsed_a) + list(parsed_b)))
    return {cls: compare_items(parsed_a.get(cls, []), parsed_b.get(cls, [])) for cls in all_cls}


# ══ Word 報告 ═════════════════════════════════════════════════════════════════

def build_word_report(name_a, name_b, comparison):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    h = doc.add_heading('商品/服務名稱比對報告', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f'檔案 A：{name_a}')
    doc.add_paragraph(f'檔案 B：{name_b}')
    doc.add_paragraph('')

    total_same    = sum(len(v["same"])    for v in comparison.values())
    total_only_a  = sum(len(v["only_a"])  for v in comparison.values())
    total_only_b  = sum(len(v["only_b"])  for v in comparison.values())
    total_similar = sum(len(v["similar"]) for v in comparison.values())

    doc.add_heading('比對摘要', level=1)
    tbl = doc.add_table(rows=2, cols=4)
    tbl.style = 'Table Grid'
    headers = ['✅ 完全相同', '⚠️ 相似（可能修改）', f'➕ 只在 {name_a}', f'➕ 只在 {name_b}']
    vals    = [str(total_same), str(total_similar), str(total_only_a), str(total_only_b)]
    for i, (h_txt, v_txt) in enumerate(zip(headers, vals)):
        tbl.rows[0].cells[i].text = h_txt
        tbl.rows[1].cells[i].text = v_txt
    doc.add_paragraph('')

    for cls, res in comparison.items():
        doc.add_heading(cls, level=1)
        has_diff = bool(res["only_a"] or res["only_b"] or res["similar"])
        if not has_diff:
            p = doc.add_paragraph(f'✅ 完全相同（{len(res["same"])} 項）')
            p.runs[0].font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        else:
            if res["similar"]:
                doc.add_heading(f'⚠️ 相似但不完全相同 — {len(res["similar"])} 項', level=2)
                t = doc.add_table(rows=1, cols=3)
                t.style = 'Table Grid'
                t.rows[0].cells[0].text = f'檔案 {name_a}'
                t.rows[0].cells[1].text = f'檔案 {name_b}'
                t.rows[0].cells[2].text = '相似度'
                for a_item, b_item, ratio in res["similar"]:
                    row = t.add_row()
                    row.cells[0].text = a_item
                    row.cells[1].text = b_item
                    row.cells[2].text = f'{int(ratio*100)}%'
                doc.add_paragraph('')
            if res["only_a"]:
                doc.add_heading(f'➕ 只在 {name_a} 有 — {len(res["only_a"])} 項', level=2)
                for item in res["only_a"]:
                    p = doc.add_paragraph(f'• {item}')
                    p.runs[0].font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            if res["only_b"]:
                doc.add_heading(f'➕ 只在 {name_b} 有 — {len(res["only_b"])} 項', level=2)
                for item in res["only_b"]:
                    p = doc.add_paragraph(f'• {item}')
                    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x5C, 0x2A)
            doc.add_paragraph(f'（相同項目：{len(res["same"])} 項）')
        doc.add_paragraph('')

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ══ UI ═══════════════════════════════════════════════════════════════════════

def input_panel(label, key_prefix):
    """回傳 (text, display_name)"""
    st.subheader(f"📄 {label}")
    display_name = st.text_input("標籤名稱", value=label, key=f"{key_prefix}_name")
    mode = st.radio("輸入方式", ["上傳檔案", "貼上文字"], key=f"{key_prefix}_mode", horizontal=True)

    text = ""
    if mode == "上傳檔案":
        f = st.file_uploader(
            "上傳 PDF 或 Word",
            type=["pdf", "docx"],
            key=f"{key_prefix}_file"
        )
        if f:
            with st.spinner("抽取文字中..."):
                try:
                    text = extract_text(f)
                    st.success(f"✅ 已抽取文字（{len(text)} 字元）")
                except Exception as e:
                    st.error(f"抽取失敗：{e}")
    else:
        text = st.text_area(
            "直接貼上文字內容",
            height=250,
            placeholder="貼上商品服務名稱，支援分號或換行分隔，例如：\nCLASS 29\nAlmond milk; Butter; Cheese\nCLASS 35\nAdvertising services; Marketing consultant",
            key=f"{key_prefix}_text"
        )

    return text, display_name


col1, col2 = st.columns(2)
with col1:
    text_a, name_a = input_panel("檔案 A", "a")
with col2:
    text_b, name_b = input_panel("檔案 B", "b")

st.divider()

if st.button("🔍 開始比對", type="primary", disabled=not (text_a and text_b)):
    with st.spinner("解析商品服務項目..."):
        parsed_a = parse_goods_services(text_a)
        parsed_b = parse_goods_services(text_b)

    with st.spinner("比對中..."):
        comparison = compare_all(parsed_a, parsed_b)

    # 摘要
    st.subheader("📊 比對摘要")
    total_same    = sum(len(v["same"])    for v in comparison.values())
    total_only_a  = sum(len(v["only_a"])  for v in comparison.values())
    total_only_b  = sum(len(v["only_b"])  for v in comparison.values())
    total_similar = sum(len(v["similar"]) for v in comparison.values())

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("✅ 完全相同",          total_same)
    m2.metric("⚠️ 相似（可能修改）",  total_similar)
    m3.metric(f"➕ 只在 {name_a} 有", total_only_a)
    m4.metric(f"➕ 只在 {name_b} 有", total_only_b)

    st.divider()

    # 逐項
    st.subheader("📋 逐項比對結果")
    for cls, res in comparison.items():
        has_diff = bool(res["only_a"] or res["only_b"] or res["similar"])
        icon = "✅" if not has_diff else "⚠️"
        with st.expander(f"{icon}  {cls}", expanded=has_diff):
            if res["similar"]:
                st.markdown(f"**⚠️ 相似但不完全相同 — {len(res['similar'])} 項**")
                for a_item, b_item, ratio in res["similar"]:
                    c1, c2 = st.columns(2)
                    c1.warning(f"**{name_a}：** {a_item}")
                    c2.warning(f"**{name_b}：** {b_item}")
                    st.caption(f"相似度 {int(ratio*100)}%")
                    st.markdown("---")
            if res["only_a"]:
                st.markdown(f"**➕ 只在 {name_a} 有（{name_b} 缺少）— {len(res['only_a'])} 項**")
                for item in res["only_a"]:
                    st.error(f"• {item}")
            if res["only_b"]:
                st.markdown(f"**➕ 只在 {name_b} 有（{name_a} 缺少）— {len(res['only_b'])} 項**")
                for item in res["only_b"]:
                    st.success(f"• {item}")
            if not has_diff:
                st.success(f"此 Class 兩份文件完全相同（{len(res['same'])} 項）")
            else:
                st.info(f"相同項目：{len(res['same'])} 項")

    st.divider()

    # 下載
    st.subheader("📥 下載報告")
    dl1, dl2 = st.columns(2)
    with dl1:
        try:
            word_bytes = build_word_report(name_a, name_b, comparison)
            st.download_button(
                label="📄 下載 Word 報告 (.docx)",
                data=word_bytes,
                file_name="comparison_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
        except Exception as e:
            st.error(f"Word 報告產生失敗：{e}")

    with dl2:
        lines = [f"比對報告\nA：{name_a}\nB：{name_b}\n" + "="*60]
        for cls, res in comparison.items():
            lines.append(f"\n【{cls}】")
            for a_i, b_i, r in res["similar"]:
                lines.append(f"  ⚠️ A：{a_i}\n     B：{b_i}  ({int(r*100)}%)")
            for i in res["only_a"]:
                lines.append(f"  ➕ 只在{name_a}：{i}")
            for i in res["only_b"]:
                lines.append(f"  ➕ 只在{name_b}：{i}")
            if not res["only_a"] and not res["only_b"] and not res["similar"]:
                lines.append("  ✅ 完全相同")
            lines.append(f"  相同：{len(res['same'])} 項")
        st.download_button(
            label="📋 下載純文字報告 (.txt)",
            data="\n".join(lines).encode("utf-8"),
            file_name="comparison_report.txt",
            mime="text/plain"
        )

elif not text_a or not text_b:
    st.info("請輸入兩份內容後點「開始比對」。")

st.divider()
st.caption("商品/服務名稱比對工具 ｜ 支援 PDF / Word / 直接貼文字")
